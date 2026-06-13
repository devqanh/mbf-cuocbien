# -*- coding: utf-8 -*-
"""Sinh file Word báo giá từ nội dung đã chốt. Chạy: python build_docx.py"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0x01, 0x53, 0xA9)
DARK  = RGBColor(0x1C, 0x27, 0x3C)
MUTED = RGBColor(0x60, 0x6A, 0x80)
FONT  = "Times New Roman"

doc = Document()

# Font mặc định
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

sec = doc.sections[0]
sec.left_margin = sec.right_margin = Cm(2.0)
sec.top_margin = sec.bottom_margin = Cm(1.8)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def run(p, text, bold=False, italic=False, color=None, size=None):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if color is not None:
        r.font.color.rgb = color
    if size is not None:
        r.font.size = Pt(size)
    return r


def heading(text, size=14, space_before=14, color=BRAND):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    run(p, text, bold=True, color=color, size=size)
    return p


def para(text="", bold=False, italic=False, color=None, align=None, size=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    if text:
        run(p, text, bold=bold, italic=italic, color=color, size=size)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    # hỗ trợ **đậm** đơn giản
    parts = text.split("**")
    for i, seg in enumerate(parts):
        if seg:
            run(p, seg, bold=(i % 2 == 1))
    return p


def make_table(headers, rows, widths=None, total_rows=None):
    total_rows = total_rows or []
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], "0153A9")
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=11)
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        is_total = ridx in total_rows
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            if i == len(row) - 1 and len(row) > 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            seg_parts = str(val).split("**")
            for j, seg in enumerate(seg_parts):
                if seg:
                    run(p, seg, bold=is_total or (j % 2 == 1), size=11)
            if is_total:
                set_cell_bg(cells[i], "E6EFFA")
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


# ====================== NỘI DUNG ======================

# Tiêu đề
para("BÁO GIÁ TRIỂN KHAI PHẦN MỀM QUẢN LÝ VẬN TẢI – LOGISTICS",
     bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, size=17, space_after=2)
para("Hệ thống quản trị điều hành Trucking trọn gói – Bảo mật chuẩn ngân hàng – Hạ tầng quốc tế tại Singapore",
     italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=10)

# Bảng thông tin
info = [
    ["Kính gửi", "CÔNG TY MBF"],
    ["Người nhận", "Ông/Bà ............................................"],
    ["Đơn vị cung cấp", "............................................"],
    ["Ngày báo giá", "...../...../20....."],
    ["Hiệu lực báo giá", "30 ngày kể từ ngày phát hành"],
    ["Mã báo giá", "BG-TRK-............"],
]
ti = doc.add_table(rows=0, cols=2)
ti.style = "Table Grid"
for k, v in info:
    c = ti.add_row().cells
    set_cell_bg(c[0], "F2F5FA")
    run(c[0].paragraphs[0], k, bold=True, color=DARK, size=11)
    run(c[1].paragraphs[0], v, bold=(k == "Kính gửi"), size=11)
ti.columns[0].width = Cm(4.5)
ti.columns[1].width = Cm(12.0)
for r in ti.rows:
    r.cells[0].width = Cm(4.5)
    r.cells[1].width = Cm(12.0)

# 1. LỜI MỞ ĐẦU
heading("1. LỜI MỞ ĐẦU")
para("Trân trọng cảm ơn Quý Công ty đã tin tưởng trao đổi nhu cầu số hoá hoạt động vận tải.")
p = para()
run(p, "Khác với các phần mềm đóng gói bán sẵn dùng chung cho mọi ngành, đây là ")
run(p, "hệ thống được xây dựng riêng (custom)", bold=True)
run(p, " theo đúng quy trình thực tế của Quý Công ty: từ quản lý lô hàng, bảng giá theo từng khách, lập bảng kê công nợ, đến tính phí xe nội bộ và quản trị đội xe. Toàn bộ dữ liệu kinh doanh quan trọng được đặt trên ")
run(p, "hạ tầng máy chủ riêng tại Singapore", bold=True)
run(p, " và bảo vệ bằng ")
run(p, "cơ chế đăng nhập 2 lớp (xác thực qua điện thoại)", bold=True)
run(p, ", tiêu chuẩn mà các ngân hàng và doanh nghiệp lớn đang áp dụng.")
p = para()
run(p, "Đây không phải là một khoản chi phí, mà là ")
run(p, "một khoản đầu tư vào tài sản số", bold=True, color=BRAND)
run(p, " vận hành lâu dài, an toàn và mở rộng được cùng sự phát triển của Quý Công ty.")

# 2. GIÁ TRỊ
heading("2. GIÁ TRỊ HỆ THỐNG MANG LẠI")
for b in [
    "**Không còn phụ thuộc file Excel rời rạc.** Mọi lô hàng, bảng giá, công nợ được lưu tập trung, không lo mất file, không lo sửa đè lên nhau.",
    "**Khép kín dòng tiền.** Lô hàng → bảng giá → bảng kê công nợ → phí xe đều liên kết, hạn chế thất thoát và cộng trùng.",
    "**Kiểm soát con người.** Phân quyền chi tiết: ai được xem, ai được sửa, ai được xoá — minh bạch và truy vết được.",
    "**An toàn tuyệt đối.** Dù lộ mật khẩu, người ngoài vẫn không thể đăng nhập nếu không có điện thoại của nhân viên duyệt.",
    "**Truy cập mọi lúc, mọi nơi.** Chạy trên trình duyệt (máy tính, điện thoại), không cần cài đặt phức tạp.",
]:
    bullet(b)

# 3. DANH MỤC TÍNH NĂNG
heading("3. DANH MỤC TÍNH NĂNG TRIỂN KHAI")
para("A. Nhóm nghiệp vụ vận tải cốt lõi", bold=True, color=DARK)
make_table(
    ["Tính năng", "Mô tả lợi ích"],
    [
        ["Quản lý Lô hàng", "Thêm/sửa/xoá lô hàng, tìm – lọc nhanh, nhập hàng loạt từ Excel và xuất Excel gửi đối tác/kế toán."],
        ["Bảng giá theo khách hàng", "Lưu bảng giá đã chốt cho từng khách; xem lại, chỉnh sửa, import — tránh tranh cãi giá về sau."],
        ["Bảng kê công nợ", "Tự động lập bảng kê cần thu và xuất Excel đúng mẫu kế toán chỉ với một cú nhấp."],
        ["Phí xe nội bộ", "Tính chi phí chuyến theo kỳ: nhiên liệu, khấu hao, phí cầu đường, lương tài… kèm cảnh báo cộng trùng."],
        ["Quản lý đội xe", "Theo dõi từng đầu xe: chi phí, khấu hao, tần suất sử dụng — biết rõ xe nào đang lãi/lỗ."],
        ["Cài đặt & danh mục", "Quản lý địa điểm, kho bãi, loại container, tài xế, cấu hình VAT/free-time… một nơi dùng chung."],
    ],
    widths=[5.0, 11.5],
)
para("", space_after=4)
para("B. Nhóm quản trị & cộng tác", bold=True, color=DARK)
make_table(
    ["Tính năng", "Mô tả lợi ích"],
    [
        ["Quản lý thành viên", "Tạo/khoá tài khoản nhân viên, đặt lại mật khẩu, kiểm soát ai đang dùng hệ thống."],
        ["Phân quyền chi tiết theo vai trò", "Cấu hình chính xác từng quyền (xem/thêm/sửa/xoá) cho từng nhóm nhân sự."],
        ["Ghi chú & Giao việc", "Giao việc kèm hạn nhắc, nhắc đến đồng nghiệp, trao đổi ngay trên từng lô hàng/báo cáo."],
        ["Thông báo tức thời (real-time)", "Nhân viên nhận thông báo ngay lập tức khi có việc mới, không cần tải lại trang."],
    ],
    widths=[5.0, 11.5],
)

# 4. HẠ TẦNG & BẢO MẬT
heading("4. HẠ TẦNG & BẢO MẬT (ĐIỂM KHÁC BIỆT CỐT LÕI)")
para("Đây là phần làm nên sự an toàn và đẳng cấp của hệ thống — yếu tố mà phần mềm giá rẻ không có.",
     italic=True, color=MUTED)
para("4.1. Máy chủ riêng đặt tại Singapore", bold=True, color=DARK)
for b in [
    "Dữ liệu đặt tại **trung tâm dữ liệu đạt chuẩn quốc tế ở Singapore** — nơi nhiều ngân hàng và tập đoàn lớn đặt hệ thống.",
    "**Tốc độ cao, ổn định**, hoạt động 24/7, chống chịu sự cố tốt hơn nhiều máy chủ trong nước giá rẻ.",
    "Được **sao lưu tự động hằng ngày** — sự cố hỏng hóc vẫn có thể khôi phục.",
    "Kết nối **mã hoá toàn trình (SSL/HTTPS)**: mọi thông tin truyền đi đều được niêm phong.",
]:
    bullet(b)
para("4.2. Đăng nhập 2 lớp – Xác thực qua điện thoại (2FA)", bold=True, color=DARK)
for b in [
    "Để vào hệ thống, nhân viên phải qua **2 lớp khoá**: (1) mật khẩu, và (2) **mã bảo mật trên điện thoại cá nhân**.",
    "Hình dung như **két sắt hai lớp khoá**: lấy được mật khẩu cũng không mở được nếu không có điện thoại của nhân viên.",
    "Mã đổi mới **mỗi 30 giây**, dùng ứng dụng chuẩn quốc tế (Google/Microsoft Authenticator) — an toàn hơn SMS.",
    "Có **bộ mã khôi phục** dự phòng khi mất điện thoại; quản trị viên **reset từ xa** — không bao giờ bị kẹt ngoài hệ thống.",
]:
    bullet(b)
para("4.3. Lớp bảo vệ dữ liệu & kiểm soát truy cập", bold=True, color=DARK)
for b in [
    "Thông tin bảo mật được **mã hoá khi lưu trữ** — kể cả người vận hành máy chủ cũng không đọc được.",
    "**Quản lý thiết bị & phiên đăng nhập**: thấy mọi thiết bị đang đăng nhập (vị trí, IP) và **đăng xuất từ xa** thiết bị lạ.",
    "Mọi thao tác quan trọng đều yêu cầu xác nhận lại — chống thao tác nhầm và truy cập trái phép.",
]:
    bullet(b)

# 5. BẢNG GIÁ TRỌN GÓI
doc.add_page_break()
heading("5. BẢNG GIÁ TRỌN GÓI")
para("Báo giá chia thành 4 hạng mục rõ ràng, mỗi hạng mục đã bao gồm trọn bộ các tính năng liên quan.",
     italic=True, color=MUTED)

def hangmuc(title, price, note, items):
    p = para(space_after=2)
    run(p, title + "  —  ", bold=True, color=DARK, size=12)
    run(p, price, bold=True, color=BRAND, size=12)
    if note:
        para(note, italic=True, color=MUTED, size=11, space_after=2)
    for it in items:
        bullet(it)
    para("", space_after=2)

hangmuc("Hạng mục 1 — Xây dựng hạ tầng lõi & Bảo mật 2 lớp", "20.000.000 VNĐ",
        "Đây là “phần móng” của toàn hệ thống — làm một lần, dùng mãi.",
        ["Nền tảng chạy trên trình duyệt (máy tính & điện thoại), không cần cài đặt.",
         "**Đăng nhập 2 lớp (2FA)** — nhân viên xác thực thêm bằng mã trên **điện thoại** mới vào được.",
         "**Mã hoá dữ liệu** & quản lý thiết bị / phiên đăng nhập (đăng xuất từ xa thiết bị lạ).",
         "**Quản lý thành viên & Phân quyền chi tiết** theo vai trò.",
         "Danh mục dùng chung (địa điểm, kho bãi, loại container, tài xế…).",
         "Ghi chú – Giao việc & **Thông báo tức thời (real-time)**."])
hangmuc("Hạng mục 2 — Quản lý Lô hàng & Bảng kê công nợ", "60.000.000 VNĐ",
        "Trái tim vận hành kinh doanh hằng ngày.",
        ["**Quản lý lô hàng**: thêm/sửa/xoá, tìm – lọc, **nhập từ Excel** và **xuất Excel**.",
         "**Bảng giá theo từng khách hàng** — lưu giá đã chốt, tránh tranh cãi về sau.",
         "**Bảng kê công nợ**: tự động lập từ lô hàng và **xuất Excel đúng mẫu kế toán hiện hành**.",
         "Liên kết khép kín **Lô hàng → Bảng giá → Bảng kê**, hạn chế thất thoát & sai sót."])
hangmuc("Hạng mục 3 — Quản lý xe & Tính chi phí từng xe", "50.000.000 VNĐ",
        "Biết chính xác mỗi đầu xe đang lãi hay lỗ.",
        ["**Quản lý đội xe**: theo dõi thông tin & tình trạng từng đầu xe.",
         "**Tính chi phí từng xe theo kỳ**: nhiên liệu, khấu hao, phí cầu đường, **lương tài xế**…",
         "**Cảnh báo cộng trùng chi phí** — không tính lặp, không thất thoát.",
         "Báo cáo **lãi/lỗ theo từng xe** để ra quyết định điều phối."])
hangmuc("Hạng mục 4 — Hạ tầng máy chủ Singapore (năm đầu)", "20.000.000 VNĐ",
        "Nơi “cất giữ” toàn bộ dữ liệu an toàn, tốc độ cao.",
        ["**Máy chủ riêng đặt tại Singapore** — trung tâm dữ liệu chuẩn quốc tế, ổn định 24/7.",
         "Tên miền + **chứng chỉ SSL/HTTPS** (mã hoá toàn trình).",
         "**Sao lưu tự động hằng ngày** — sự cố vẫn khôi phục được."])

heading("TỔNG GIÁ TRỊ HỢP ĐỒNG (NĂM ĐẦU)", size=13, space_before=8)
make_table(
    ["Hạng mục", "Thành tiền (VNĐ)"],
    [
        ["1. Xây dựng hạ tầng lõi & Bảo mật 2 lớp", "20.000.000"],
        ["2. Quản lý Lô hàng & Bảng kê công nợ", "60.000.000"],
        ["3. Quản lý xe & Tính chi phí từng xe (gồm lương tài xế)", "50.000.000"],
        ["4. Hạ tầng máy chủ Singapore (năm đầu)", "20.000.000"],
        ["**Cộng (chưa bao gồm VAT)**", "**150.000.000**"],
        ["Thuế GTGT (VAT 10%)", "15.000.000"],
        ["**TỔNG THANH TOÁN (đã bao gồm VAT)**", "**165.000.000**"],
    ],
    widths=[11.5, 5.0],
    total_rows=[4, 6],
)
p = para(space_after=2)
run(p, "Bằng chữ: ", bold=True)
run(p, "Một trăm sáu mươi lăm triệu đồng chẵn.", italic=True)
para("Giá trị hợp đồng chưa bao gồm VAT là 150.000.000 VNĐ; thuế GTGT 10% được tính riêng theo quy định.",
     italic=True, color=MUTED, size=11)

heading("Chi phí từ năm thứ 2 trở đi", size=13)
make_table(
    ["Hạng mục", "Chi phí"],
    [["Phí hạ tầng máy chủ Singapore (máy chủ riêng, tên miền, SSL, sao lưu tự động, duy trì vận hành)",
      "**20.000.000 / năm** (chưa VAT)"]],
    widths=[11.5, 5.0],
)
p = para(space_after=4)
run(p, "Hệ thống sau khi hoàn thiện và chạy ổn định ở năm đầu thì các năm sau hầu như không phát sinh lỗi; "
        "chi phí duy trì chỉ là 20 triệu/năm tiền hạ tầng. Các ", italic=True, color=MUTED, size=11)
run(p, "tính năng phát sinh mới", italic=True, bold=True, color=MUTED, size=11)
run(p, " (nếu yêu cầu thêm) sẽ được báo giá & tính phí riêng theo từng hạng mục.",
    italic=True, color=MUTED, size=11)

# 6. THANH TOÁN
heading("6. ĐIỀU KHOẢN THANH TOÁN")
para("Thanh toán 3 đợt theo tiến độ (tính trên tổng đã bao gồm VAT: 165.000.000 VNĐ):", size=11)
make_table(
    ["Đợt", "Thời điểm", "Tỷ lệ", "Số tiền (VNĐ)"],
    [
        ["Đợt 1", "Ngay khi ký hợp đồng", "40%", "66.000.000"],
        ["Đợt 2", "Khi bàn giao bản chạy thử (demo)", "40%", "66.000.000"],
        ["Đợt 3", "Sau nghiệm thu & bàn giao chính thức", "20%", "33.000.000"],
    ],
    widths=[2.0, 8.5, 2.0, 4.0],
)

# 7. TIẾN ĐỘ
heading("7. TIẾN ĐỘ & BÀN GIAO")
make_table(
    ["Giai đoạn", "Thời gian dự kiến"],
    [
        ["Khảo sát & thống nhất yêu cầu", "1 tuần"],
        ["Phát triển & hoàn thiện hệ thống", "4 – 6 tuần"],
        ["Cài đặt hạ tầng Singapore + bảo mật", "song song"],
        ["Đào tạo, nghiệm thu & bàn giao", "1 tuần"],
        ["**Tổng thời gian**", "**≈ 6 – 8 tuần**"],
    ],
    widths=[11.5, 5.0],
    total_rows=[4],
)

# 8. CAM KẾT
heading("8. CHÚNG TÔI CAM KẾT")
for b in [
    "**Bảo hành 1 năm** kể từ ngày bàn giao — sửa lỗi miễn phí, hỗ trợ kỹ thuật trong giờ hành chính (áp dụng khi Quý Công ty duy trì sử dụng dịch vụ trong năm đầu).",
    "**Tính năng phát sinh mới** sẽ được **báo giá & tính phí riêng** theo từng hạng mục — minh bạch, thống nhất trước khi làm.",
    "**Dữ liệu là của Quý Công ty** — có thể xuất ra bất cứ lúc nào, không bị “giam” dữ liệu.",
    "**Đào tạo đến khi nhân viên sử dụng thành thạo**, kèm tài liệu hướng dẫn.",
    "**Bảo mật xuyên suốt**: hạ tầng quốc tế + đăng nhập 2 lớp + mã hoá dữ liệu.",
]:
    bullet(b)

para("", space_after=8)
para("Hệ thống tốt không chỉ giúp Quý Công ty làm việc nhanh hơn, mà còn bảo vệ tài sản dữ liệu — "
     "thứ ngày càng quý giá trong kinh doanh. Chúng tôi rất mong được đồng hành cùng Quý Công ty "
     "trên hành trình số hoá an toàn và bền vững.", italic=True, color=BRAND)
para("Trân trọng cảm ơn!", bold=True)
para("", space_after=4)
para("Mọi trao đổi chi tiết, xin liên hệ:", italic=True, color=MUTED, size=11)
p = para(size=11)
run(p, "............................", bold=True)
run(p, "   —   ĐT/Zalo: ........................   —   Email: ........................")

out = r"c:\laragon\www\cuocbien\dev\bao-gia.docx"
doc.save(out)
print("Saved:", out)
